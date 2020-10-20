from docx import Document


class WordTemplate(object):

    def __init__(self, filename):
        self.document = Document(filename)

    def save(self, filename):
        self.document.save(filename)

    def change_paragraph(self, paragraph, filter_format):
        for sformat, content in filter_format.items():
            search_word = "{" + sformat + "}"
            needCheck = True if search_word in paragraph.text else False
            for run in paragraph.runs:
                if search_word in run.text:
                    replaced = run.text.replace(search_word, content)
                    run = run.clear()
                    run.add_text(replaced)
                    needCheck = False
            if needCheck:
                start_idx = 0
                end_idx = 0
                complete_word = ""
                i = 0
                broken_keys = []
                for run in paragraph.runs:
                    if run.text == "{":
                        start_idx = i
                        complete_word = run.text
                    else:
                        end_idx = i
                        complete_word = complete_word + run.text
                        if run.text == "}":
                            replaced = complete_word.replace(search_word, content)
                            broken_keys.append((start_idx, end_idx, replaced))
                    i = i + 1

                for (start_idx, end_idx, replaced) in broken_keys:
                    for run_idx, run in enumerate(paragraph.runs):
                        if start_idx <= run_idx and run_idx <= end_idx:
                            run = run.clear()
                        if run_idx == end_idx:
                            run.add_text(replaced)

    def replace(self, filter_format):
        """ code is very rubbish, but it Cross-platform """
        for paragraph in self.document.paragraphs:
            self.change_paragraph(paragraph, filter_format)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.change_paragraph(paragraph, filter_format)
